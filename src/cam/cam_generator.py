"""
cam_generator.py — Credit Appraisal Memo (CAM) document generator.

Produces a professionally formatted Microsoft Word (.docx) document
containing all standard sections of an Indian bank CAM.

Public API
----------
    from src.cam.cam_generator import CAMGenerator

    gen = CAMGenerator()
    path = gen.generate_cam(
        company_data    = {...},
        scoring_result  = {...},   # from CreditScorer.score()
        research_report = {...},   # from ResearchAgent / SynthesizerAgent
        five_cs_text    = {...},   # from FiveCsWriter.write_* methods
        output_path     = "outputs/ACME_CAM_2026.docx",
    )
    print("CAM saved to", path)

Document sections
-----------------
1. Cover Page             — name, loan amount, date, CAM ref, CONFIDENTIAL
2. Executive Summary      — 2-col table with key decision metrics
3. Company Background     — incorporation, CIN, directors, description
4. Financial Analysis     — 3-year ratio table with trend arrows
5. GST & Bank Recon       — key findings, red-flagged items highlighted
6. Five Cs                — LLM-written narrative per C
7. Risk Score             — score/band + SHAP top factors table
8. Recommendation         — amount, rate, tenure, rationale (coloured box)
9. Early Warning Indicators — monitoring trigger bullet list
"""

from __future__ import annotations

import logging
import os
import sys
import uuid
from datetime import date, datetime
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Project-root path bootstrap
# ---------------------------------------------------------------------------
_PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

from dotenv import load_dotenv  # noqa: E402
load_dotenv(_PROJECT_ROOT / ".env")

# ---------------------------------------------------------------------------
# python-docx imports (package ships as 'docx', installed as 'python-docx')
# ---------------------------------------------------------------------------
from docx import Document                                             # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT                       # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING       # noqa: E402
from docx.oxml import OxmlElement                                     # noqa: E402
from docx.oxml.ns import qn                                           # noqa: E402
from docx.shared import Inches, Pt, RGBColor                         # noqa: E402

logger = logging.getLogger("intelli_credit.cam.cam_generator")

# ---------------------------------------------------------------------------
# Colour palette (RGB tuples)
# ---------------------------------------------------------------------------
# Navy header background
_NAVY   = RGBColor(0x1F, 0x36, 0x64)       # #1F3664
# Alternate table row (light gray)
_GRAY   = RGBColor(0xD9, 0xD9, 0xD9)       # #D9D9D9
# Red flag text / cell shading
_RED    = RGBColor(0xC0, 0x00, 0x00)       # #C00000
# Green positive signal text
_GREEN  = RGBColor(0x37, 0x86, 0x30)       # #378630
# Amber / conditional decision
_AMBER  = RGBColor(0xBF, 0x81, 0x00)       # #BF8100
# White (header text)
_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
# Dark text
_BLACK  = RGBColor(0x00, 0x00, 0x00)
# Recommendation box background (light blue)
_LIGHT_BLUE = RGBColor(0xDE, 0xEB, 0xF7)  # #DEEBF7

# ---------------------------------------------------------------------------
# Risk-band → colour mapping
# ---------------------------------------------------------------------------
_BAND_COLOUR: dict[str, RGBColor] = {
    "HIGH":   _RED,
    "MEDIUM": _AMBER,
    "LOW":    _GREEN,
    "PRIME":  RGBColor(0x00, 0x70, 0xC0),  # blue
}

# Decision keyword → colour
_DECISION_COLOUR: dict[str, RGBColor] = {
    "APPROVE":      _GREEN,
    "REJECT":       _RED,
    "CONDITIONAL":  _AMBER,
}

# ---------------------------------------------------------------------------
# Default EWI triggers (used when caller does not supply them)
# ---------------------------------------------------------------------------
_DEFAULT_EWI: list[str] = [
    "DSCR falls below 1.10 for two consecutive quarters",
    "Current Ratio drops below 0.90",
    "Any new eCourts filing classified as CRITICAL or INSOLVENCY severity",
    "GST filing non-compliance for more than one period",
    "ITC gap (GSTR-2A mismatch) exceeds 20% in any quarter",
    "Bounce of cheque / ECS mandate on loan account",
    "Director listed as wilful defaulter by RBI",
    "Adverse media coverage classified HIGH-RISK by news intelligence module",
    "Revenue decline >15% YoY as per quarterly MIS",
    "Any NCLT / IBC proceedings initiated against the company",
]


# ===========================================================================
# Low-level helpers
# ===========================================================================

def _hex_to_rgb_str(colour: RGBColor) -> str:
    """Return 6-digit hex string for OOXML shading (e.g. '1F3664').

    RGBColor subclasses tuple in python-docx 1.2.x — access by index.
    """
    return f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}"


def _set_cell_bg(cell: Any, colour: RGBColor) -> None:
    """Fill a table cell background with *colour* using raw OOXML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _hex_to_rgb_str(colour))
    tcPr.append(shd)


def _set_row_bg(row: Any, colour: RGBColor) -> None:
    for cell in row.cells:
        _set_cell_bg(cell, colour)


def _cell_text(
    cell:      Any,
    text:      str,
    bold:      bool      = False,
    italic:    bool      = False,
    colour:    RGBColor | None = None,
    font_size: int       = 10,
    align:     int       = WD_ALIGN_PARAGRAPH.LEFT,
    wrap:      bool      = True,
) -> None:
    """
    Replace *cell* content with a single paragraph run carrying the given style.
    Clears any existing content first.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(str(text) if text is not None else "")
    run.bold    = bold
    run.italic  = italic
    run.font.size = Pt(font_size)
    if colour:
        run.font.color.rgb = colour


def _add_navy_heading(doc: Any, text: str, level: int = 1) -> Any:
    """Add a section heading with navy colour and white-ish uppercase text."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run  = para.add_run(text.upper())
    run.bold = True
    run.font.size   = Pt(13 if level == 1 else 11)
    run.font.color.rgb = _NAVY
    return para


def _add_subheading(doc: Any, text: str) -> Any:
    """Add a subsection heading (navy, slightly smaller)."""
    return _add_navy_heading(doc, text, level=2)


def _make_table(
    doc:       Any,
    rows:      int,
    cols:      int,
    col_widths: list[float] | None = None,   # inches
) -> Any:
    """Create a borderless-style table and optionally set column widths."""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    return tbl


def _page_break(doc: Any) -> None:
    doc.add_page_break()


def _spacer(doc: Any, space_pt: int = 6) -> None:
    """Add an empty paragraph as a vertical spacer."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(space_pt)


def _trend_arrow(current: float | None, previous: float | None) -> str:
    """Return ↑ / ↓ / → depending on change between two values."""
    if current is None or previous is None:
        return "—"
    if current > previous * 1.02:
        return "↑"
    if current < previous * 0.98:
        return "↓"
    return "→"


def _fmt(value: Any, prefix: str = "", suffix: str = "", na: str = "N/A") -> str:
    if value is None:
        return na
    try:
        f = float(value)
        return f"{prefix}{f:,.2f}{suffix}"
    except (TypeError, ValueError):
        return f"{prefix}{value}{suffix}"


def _safe_str(value: Any, na: str = "N/A") -> str:
    if value is None or (isinstance(value, float) and value != value):
        return na
    return str(value).strip() or na


# ===========================================================================
# CAMGenerator
# ===========================================================================

class CAMGenerator:
    """
    Generates a fully-formatted Credit Appraisal Memo as a .docx file.

    Parameters
    ----------
    page_width_inches : float
        Usable page width (default A4 portrait with 1-inch margins = 6.27 in).
    """

    def __init__(self, page_width_inches: float = 6.27) -> None:
        self._pw = page_width_inches

    # ==================================================================
    # Main entry point
    # ==================================================================

    def generate_cam(
        self,
        company_data:    dict[str, Any],
        scoring_result:  dict[str, Any],
        research_report: dict[str, Any],
        five_cs_text:    dict[str, Any],
        output_path:     str | Path,
    ) -> Path:
        """
        Assemble all CAM sections into a Word document and save it.

        Parameters
        ----------
        company_data : dict
            Company & loan metadata.  Expected keys (all optional with defaults):

            ======================  ============================================
            name                    Legal entity name
            cin                     CIN / registration number
            incorporation_date      Date string or date object
            directors               list[str | dict] — director names / details
            business_description    Short business overview paragraph
            loan_amount_requested   str or numeric (₹)
            recommended_amount      str or numeric (₹)
            interest_rate           str (e.g. "12.5% p.a.")
            tenure                  str (e.g. "60 months")
            decision                "APPROVE" | "REJECT" | "CONDITIONAL"
            decision_rationale      str — 3-sentence rationale
            financials_3yr          list[dict] — one dict per year with keys:
                                     year, revenue, ebitda, pat, de_ratio,
                                     current_ratio, dscr
            gst_findings            dict — from EWS / GST reconciler
            bank_findings           dict — from BankStatementAnalyzer
            ews_flags               dict — {flag_name: level}
            ewi_triggers            list[str] — custom monitoring triggers
            ======================  ============================================

        scoring_result : dict
            From ``CreditScorer.score()``.  Keys:
            ``default_probability``, ``risk_score``, ``risk_band``,
            ``shap_explanations`` → {top_risk_factors, top_positive_factors}.

        research_report : dict
            From ``ResearchAgent`` / ``SynthesizerAgent``.  Keys consumed:
            ``news_summary``, ``promoter_risk_flag``,
            ``regulatory_compliance_summary``, ``key_red_flags``,
            ``positive_signals``.

        five_cs_text : dict
            Keys: ``CHARACTER``, ``CAPACITY``, ``CAPITAL``,
            ``COLLATERAL``, ``CONDITIONS``.  Each value is either a plain
            string or the dict returned by ``FiveCsWriter.write_*()``.

        output_path : str | Path
            Where to save the .docx.  Parent directories are created.

        Returns
        -------
        Path — absolute path to the saved document.
        """
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # ── Global page margins (1 inch all sides) ─────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

        # ── Set default paragraph font ──────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # ── Build each section ──────────────────────────────────────────
        cam_ref = company_data.get("cam_ref") or f"CAM/{date.today().strftime('%Y%m')}/{uuid.uuid4().hex[:6].upper()}"

        self._cover_page(doc, company_data, cam_ref)
        _page_break(doc)

        self._executive_summary(doc, company_data, scoring_result)
        _page_break(doc)

        self._company_background(doc, company_data, research_report)
        _page_break(doc)

        self._financial_analysis(doc, company_data)
        _page_break(doc)

        self._gst_bank_recon(doc, company_data, research_report)
        _page_break(doc)

        self._five_cs(doc, five_cs_text)
        _page_break(doc)

        self._risk_score_section(doc, scoring_result)
        _page_break(doc)

        self._recommendation(doc, company_data, scoring_result)
        _page_break(doc)

        self._early_warning_indicators(doc, company_data)

        doc.save(str(out))
        logger.info("CAM saved → %s", out.resolve())
        return out.resolve()

    # ==================================================================
    # Section 1 — Cover Page
    # ==================================================================

    def _cover_page(
        self,
        doc:          Any,
        company_data: dict[str, Any],
        cam_ref:      str,
    ) -> None:
        name        = _safe_str(company_data.get("name"), "Borrower Name")
        loan_amt    = _safe_str(company_data.get("loan_amount_requested"), "To be specified")
        today       = date.today().strftime("%d %B %Y")

        # STRICTLY CONFIDENTIAL band (navy box at top)
        conf_para = doc.add_paragraph()
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_run  = conf_para.add_run("STRICTLY CONFIDENTIAL")
        conf_run.bold           = True
        conf_run.font.size      = Pt(11)
        conf_run.font.color.rgb = _WHITE
        # shade paragraph via XML (paragraph background not natively supported;
        # we use a single-cell table as a coloured banner instead)
        # Replace above paragraph with a 1-cell table banner:
        conf_para.clear()
        banner_tbl = _make_table(doc, 1, 1, col_widths=[self._pw])
        cell = banner_tbl.rows[0].cells[0]
        _set_cell_bg(cell, _NAVY)
        _cell_text(cell, "STRICTLY CONFIDENTIAL", bold=True, font_size=11,
                   colour=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        # Remove the empty paragraph left before the table
        conf_para._element.getparent().remove(conf_para._element)

        _spacer(doc, 18)

        # Institution name
        inst_para = doc.add_paragraph()
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inst_run  = inst_para.add_run("CREDIT APPRAISAL MEMORANDUM")
        inst_run.bold           = True
        inst_run.font.size      = Pt(18)
        inst_run.font.color.rgb = _NAVY

        _spacer(doc, 6)

        # Divider
        div_para = doc.add_paragraph("─" * 72)
        div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _spacer(doc, 18)

        # Company name (large)
        co_para = doc.add_paragraph()
        co_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        co_run  = co_para.add_run(name)
        co_run.bold           = True
        co_run.font.size      = Pt(20)
        co_run.font.color.rgb = _NAVY

        _spacer(doc, 12)

        # Loan amount
        la_para = doc.add_paragraph()
        la_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        la_run  = la_para.add_run(f"Proposed Credit Facility: {loan_amt}")
        la_run.font.size = Pt(13)
        la_run.bold = True

        _spacer(doc, 40)

        # Meta table (date / CAM ref)
        meta = _make_table(doc, 2, 2, col_widths=[3.0, 3.27])
        rows_data = [
            ("Date of Appraisal:", today),
            ("CAM Reference No.:", cam_ref),
        ]
        for i, (label, val) in enumerate(rows_data):
            _cell_text(meta.rows[i].cells[0], label, bold=True, font_size=10)
            _cell_text(meta.rows[i].cells[1], val,   bold=False, font_size=10)

        _spacer(doc, 40)

        # Footer note
        foot_para = doc.add_paragraph()
        foot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        foot_run  = foot_para.add_run(
            "This document is prepared for internal credit committee use only. "
            "Unauthorised distribution is strictly prohibited."
        )
        foot_run.italic    = True
        foot_run.font.size = Pt(9)
        foot_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # ==================================================================
    # Section 2 — Executive Summary
    # ==================================================================

    def _executive_summary(
        self,
        doc:            Any,
        company_data:   dict[str, Any],
        scoring_result: dict[str, Any],
    ) -> None:
        _add_navy_heading(doc, "Executive Summary")
        _spacer(doc)

        name        = _safe_str(company_data.get("name"))
        loan_req    = _safe_str(company_data.get("loan_amount_requested"))
        rec_amt     = _safe_str(company_data.get("recommended_amount"))
        rate        = _safe_str(company_data.get("interest_rate"))
        tenure      = _safe_str(company_data.get("tenure"))
        decision    = _safe_str(company_data.get("decision"), "PENDING").upper()
        risk_score  = scoring_result.get("risk_score")
        risk_band   = _safe_str(scoring_result.get("risk_band"))
        default_prob= scoring_result.get("default_probability")

        rows = [
            ("Applicant",               name),
            ("Loan Amount Requested",   loan_req),
            ("Recommended Amount",      rec_amt),
            ("Interest Rate Proposed",  rate),
            ("Tenure",                  tenure),
            ("Risk Score (0–10)",       f"{risk_score:.2f} / 10.00" if risk_score is not None else "N/A"),
            ("Default Probability",     f"{default_prob:.2%}" if default_prob is not None else "N/A"),
            ("Risk Band",               risk_band),
            ("Decision",                decision),
        ]

        tbl = _make_table(doc, len(rows), 2, col_widths=[2.3, self._pw - 2.3])
        for i, (label, value) in enumerate(rows):
            bg = _GRAY if i % 2 == 0 else None

            label_cell = tbl.rows[i].cells[0]
            value_cell = tbl.rows[i].cells[1]

            if bg:
                _set_cell_bg(label_cell, bg)
                _set_cell_bg(value_cell, bg)

            _cell_text(label_cell, label, bold=True, font_size=10)

            # Special colouring for Risk Band and Decision
            if label == "Risk Band":
                colour = _BAND_COLOUR.get(risk_band, _BLACK)
                _cell_text(value_cell, value, bold=True, colour=colour, font_size=10)
            elif label == "Decision":
                # Find matching keyword in decision string
                col = _BLACK
                for kw, c in _DECISION_COLOUR.items():
                    if decision.startswith(kw):
                        col = c
                        break
                _cell_text(value_cell, value, bold=True, colour=col, font_size=11)
            else:
                _cell_text(value_cell, value, font_size=10)

        _spacer(doc)

        # Header row shade
        _set_row_bg(tbl.rows[0], _NAVY)
        _cell_text(tbl.rows[0].cells[0], rows[0][0], bold=True, colour=_WHITE, font_size=10)
        _cell_text(tbl.rows[0].cells[1], rows[0][1], bold=True, font_size=10)

    # ==================================================================
    # Section 3 — Company Background
    # ==================================================================

    def _company_background(
        self,
        doc:             Any,
        company_data:    dict[str, Any],
        research_report: dict[str, Any],
    ) -> None:
        _add_navy_heading(doc, "Company Background")
        _spacer(doc)

        cin     = _safe_str(company_data.get("cin"))
        inc_dt  = _safe_str(company_data.get("incorporation_date"))
        biz_desc= _safe_str(company_data.get("business_description"), "Business description not provided.")

        # Key facts table
        facts = [
            ("Company Name",       _safe_str(company_data.get("name"))),
            ("CIN / Reg. Number",  cin),
            ("Date of Incorporation", inc_dt),
            ("Registered Office",  _safe_str(company_data.get("registered_office"))),
            ("Line of Business",   _safe_str(company_data.get("business_sector"))),
            ("Constitution",       _safe_str(company_data.get("constitution"), "Private Limited Company")),
        ]

        tbl = _make_table(doc, len(facts), 2, col_widths=[2.3, self._pw - 2.3])
        for i, (lbl, val) in enumerate(facts):
            if i == 0:
                _set_row_bg(tbl.rows[i], _NAVY)
                _cell_text(tbl.rows[i].cells[0], lbl, bold=True, colour=_WHITE, font_size=10)
                _cell_text(tbl.rows[i].cells[1], val, bold=True, colour=_WHITE, font_size=10)
            else:
                bg = _GRAY if i % 2 == 0 else None
                if bg:
                    _set_row_bg(tbl.rows[i], bg)
                _cell_text(tbl.rows[i].cells[0], lbl, bold=True, font_size=10)
                _cell_text(tbl.rows[i].cells[1], val,  font_size=10)

        _spacer(doc, 10)
        _add_subheading(doc, "Directors / Key Management Personnel")
        _spacer(doc)

        directors = company_data.get("directors") or []
        if not directors:
            directors = [{"name": "Not provided", "din": "—", "designation": "—"}]

        # Normalise: accept list[str] or list[dict]
        dir_dicts: list[dict[str, str]] = []
        for d in directors:
            if isinstance(d, str):
                dir_dicts.append({"name": d, "din": "—", "designation": "Director"})
            elif isinstance(d, dict):
                dir_dicts.append({
                    "name":        d.get("name", "—"),
                    "din":         d.get("din", d.get("DIN", "—")),
                    "designation": d.get("designation", d.get("role", "Director")),
                })

        dir_tbl = _make_table(doc, len(dir_dicts) + 1, 3,
                              col_widths=[2.8, 1.3, 2.17])
        # Header
        _set_row_bg(dir_tbl.rows[0], _NAVY)
        for col, hdr in enumerate(["Director Name", "DIN", "Designation"]):
            _cell_text(dir_tbl.rows[0].cells[col], hdr,
                       bold=True, colour=_WHITE, font_size=10)
        # Data
        for i, d in enumerate(dir_dicts, start=1):
            bg = _GRAY if i % 2 == 0 else None
            if bg:
                _set_row_bg(dir_tbl.rows[i], bg)
            _cell_text(dir_tbl.rows[i].cells[0], d["name"],        font_size=10)
            _cell_text(dir_tbl.rows[i].cells[1], d["din"],         font_size=10)
            _cell_text(dir_tbl.rows[i].cells[2], d["designation"], font_size=10)

        _spacer(doc, 10)
        _add_subheading(doc, "Business Description")
        _spacer(doc)

        biz_para = doc.add_paragraph(biz_desc)
        biz_para.paragraph_format.space_after = Pt(6)

        # Key red flags from research (if any)
        _synth = research_report.get("synthesis_report")
        red_flags = (
            research_report.get("key_red_flags")
            or (isinstance(_synth, dict) and _synth.get("key_red_flags"))
            or []
        )
        if red_flags:
            _spacer(doc)
            _add_subheading(doc, "Key Red Flags (External Intelligence)")
            _spacer(doc)
            for flag in red_flags:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(str(flag))
                run.font.color.rgb = _RED
                run.font.size      = Pt(10)

    # ==================================================================
    # Section 4 — Financial Analysis (3-year ratio table)
    # ==================================================================

    def _financial_analysis(
        self,
        doc:          Any,
        company_data: dict[str, Any],
    ) -> None:
        _add_navy_heading(doc, "Financial Analysis")
        _spacer(doc)

        # financials_3yr: list of dicts, most-recent first or last
        fin_data: list[dict[str, Any]] = company_data.get("financials_3yr") or []

        # Ensure exactly 3 entries (pad with empty dicts if fewer)
        while len(fin_data) < 3:
            fin_data.insert(0, {})
        fin_data = fin_data[-3:]   # take most recent 3

        years = [d.get("year", f"FY{i+1}") for i, d in enumerate(fin_data)]

        metrics = [
            ("Revenue (₹ Cr)",      "revenue"),
            ("EBITDA (₹ Cr)",       "ebitda"),
            ("PAT (₹ Cr)",          "pat"),
            ("EBITDA Margin (%)",    "ebitda_margin_pct"),
            ("PAT Margin (%)",       "pat_margin_pct"),
            ("Debt / Equity",        "de_ratio"),
            ("Current Ratio",        "current_ratio"),
            ("DSCR",                 "dscr"),
            ("Revenue Growth (%)",   "revenue_growth_pct"),
        ]

        # Header: Metric | FY1 | FY2 | FY3 | Trend
        col_w = [2.3] + [(self._pw - 2.3 - 0.7) / 3] * 3 + [0.7]
        tbl = _make_table(doc, len(metrics) + 1, 5, col_widths=col_w)

        # Header row
        _set_row_bg(tbl.rows[0], _NAVY)
        for j, hdr in enumerate(["Metric"] + years + ["Trend"]):
            _cell_text(tbl.rows[0].cells[j], hdr,
                       bold=True, colour=_WHITE, font_size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT)

        for i, (label, key) in enumerate(metrics):
            row = tbl.rows[i + 1]
            bg = _GRAY if i % 2 == 0 else None
            if bg:
                _set_row_bg(row, bg)

            vals = [d.get(key) for d in fin_data]
            arrow = _trend_arrow(vals[-1], vals[-2])

            # Colour bad trends red, good green (context-dependent heuristic)
            _downward_bad = {"revenue", "ebitda", "pat", "ebitda_margin_pct",
                             "pat_margin_pct", "current_ratio", "dscr",
                             "revenue_growth_pct"}
            _upward_bad   = {"de_ratio"}
            arrow_colour  = _BLACK
            if key in _downward_bad:
                arrow_colour = _GREEN if arrow == "↑" else (_RED if arrow == "↓" else _BLACK)
            elif key in _upward_bad:
                arrow_colour = _RED if arrow == "↑" else (_GREEN if arrow == "↓" else _BLACK)

            _cell_text(row.cells[0], label, bold=True, font_size=10)
            for j, v in enumerate(vals):
                _cell_text(row.cells[j + 1], _fmt(v), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row.cells[4], arrow, bold=True, font_size=11,
                       colour=arrow_colour, align=WD_ALIGN_PARAGRAPH.CENTER)

        _spacer(doc)

        note = doc.add_paragraph("* Trend arrow reflects change from penultimate to latest year.  "
                                 "↑ Improvement   ↓ Deterioration   → Stable.")
        note.runs[0].font.size  = Pt(8)
        note.runs[0].italic     = True
        note.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # ==================================================================
    # Section 5 — GST & Bank Reconciliation
    # ==================================================================

    def _gst_bank_recon(
        self,
        doc:             Any,
        company_data:    dict[str, Any],
        research_report: dict[str, Any],
    ) -> None:
        _add_navy_heading(doc, "GST & Bank Statement Reconciliation")
        _spacer(doc)

        gst  = company_data.get("gst_findings")  or {}
        bank = company_data.get("bank_findings") or {}
        ews  = company_data.get("ews_flags")      or {}

        # ── GST findings ───────────────────────────────────────────────
        _add_subheading(doc, "GST Analysis")
        _spacer(doc)

        gst_rows = [
            ("GST Health Score",         _safe_str(gst.get("health_score") or gst.get("gst_health_score"))),
            ("GST Grade",                _safe_str(gst.get("grade"))),
            ("ITC Gap (GSTR-2A)",        _safe_str(gst.get("itc_gap_pct") or gst.get("itc_overall_risk"))),
            ("Turnover Consistency",     _safe_str(gst.get("turnover_consistency") or gst.get("turnover_flag"))),
            ("Filing Regularity",        _safe_str(gst.get("filing_regularity"))),
            ("Fictitious Vendor Count",  _safe_str(gst.get("fictitious_vendors", 0))),
            ("Circular Trading Flag",    _safe_str(ews.get("circular_trading_risk") or gst.get("circular_trading_confidence"))),
            ("GST ITC Fraud Risk",       _safe_str(ews.get("gst_itc_fraud_risk"))),
        ]

        gst_tbl = _make_table(doc, len(gst_rows) + 1, 2, col_widths=[2.8, self._pw - 2.8])
        _set_row_bg(gst_tbl.rows[0], _NAVY)
        _cell_text(gst_tbl.rows[0].cells[0], "GST Metric",  bold=True, colour=_WHITE, font_size=10)
        _cell_text(gst_tbl.rows[0].cells[1], "Finding",     bold=True, colour=_WHITE, font_size=10)

        _RED_FLAG_VALS = {"HIGH", "HIGH_RISK", "RED", "CRITICAL", "FRAUD", "SMA-2"}
        for i, (lbl, val) in enumerate(gst_rows):
            row = gst_tbl.rows[i + 1]
            bg  = _GRAY if i % 2 == 0 else None
            if bg:
                _set_row_bg(row, bg)
            _cell_text(row.cells[0], lbl, bold=True, font_size=10)
            is_red = any(r in val.upper() for r in _RED_FLAG_VALS)
            _cell_text(row.cells[1], val, font_size=10,
                       colour=_RED if is_red else _BLACK,
                       bold=is_red)

        _spacer(doc, 10)

        # ── Bank statement findings ────────────────────────────────────
        _add_subheading(doc, "Bank Statement Analysis")
        _spacer(doc)

        bank_rows = [
            ("Average Monthly Balance",  _safe_str(bank.get("avg_monthly_balance"))),
            ("Debit / Credit Ratio",     _safe_str(bank.get("debit_credit_ratio"))),
            ("Cheque / ECS Bounce Count",_safe_str(bank.get("bounce_count", 0))),
            ("UPI Concentration (%)",    _safe_str(bank.get("upi_concentration"))),
            ("Cash Stress Flag",         _safe_str(bank.get("cash_stress_flag") or ews.get("cash_stress_risk"))),
            ("Revenue Inflation Flag",   _safe_str(bank.get("revenue_inflation_flag") or ews.get("revenue_inflation_risk"))),
        ]

        bank_tbl = _make_table(doc, len(bank_rows) + 1, 2, col_widths=[2.8, self._pw - 2.8])
        _set_row_bg(bank_tbl.rows[0], _NAVY)
        _cell_text(bank_tbl.rows[0].cells[0], "Bank Metric", bold=True, colour=_WHITE, font_size=10)
        _cell_text(bank_tbl.rows[0].cells[1], "Finding",     bold=True, colour=_WHITE, font_size=10)

        for i, (lbl, val) in enumerate(bank_rows):
            row = bank_tbl.rows[i + 1]
            bg  = _GRAY if i % 2 == 0 else None
            if bg:
                _set_row_bg(row, bg)
            _cell_text(row.cells[0], lbl, bold=True, font_size=10)
            is_red = any(r in val.upper() for r in _RED_FLAG_VALS) or (
                lbl == "Cheque / ECS Bounce Count" and _safe_str(val, "0") not in ("0", "N/A")
            )
            _cell_text(row.cells[1], val, font_size=10,
                       colour=_RED if is_red else _BLACK,
                       bold=is_red)

        _spacer(doc, 10)

        # ── EWS flags table ────────────────────────────────────────────
        if ews:
            _add_subheading(doc, "Early Warning System (EWS) Flags")
            _spacer(doc)

            ews_tbl = _make_table(doc, len(ews) + 1, 2, col_widths=[3.5, self._pw - 3.5])
            _set_row_bg(ews_tbl.rows[0], _NAVY)
            _cell_text(ews_tbl.rows[0].cells[0], "EWS Flag",  bold=True, colour=_WHITE, font_size=10)
            _cell_text(ews_tbl.rows[0].cells[1], "Level",     bold=True, colour=_WHITE, font_size=10)

            flag_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2, "CLEAR": 3}
            sorted_flags = sorted(ews.items(), key=lambda x: flag_order.get(x[1].upper(), 4))

            for i, (flag_name, level) in enumerate(sorted_flags):
                row = ews_tbl.rows[i + 1]
                bg  = _GRAY if i % 2 == 0 else None
                if bg:
                    _set_row_bg(row, bg)
                label_str = flag_name.replace("_", " ").title()
                colour    = _BAND_COLOUR.get(level.upper(), _BLACK)
                _cell_text(row.cells[0], label_str,  bold=True,  font_size=10)
                _cell_text(row.cells[1], level.upper(), bold=True, colour=colour, font_size=10)

    # ==================================================================
    # Section 6 — Five Cs of Credit
    # ==================================================================

    def _five_cs(
        self,
        doc:          Any,
        five_cs_text: dict[str, Any],
    ) -> None:
        _add_navy_heading(doc, "Credit Assessment — The Five C's")
        _spacer(doc)

        section_labels = {
            "CHARACTER":  "1. Character — Management Integrity & Track Record",
            "CAPACITY":   "2. Capacity — Debt Servicing Ability",
            "CAPITAL":    "3. Capital — Net Worth & Equity Cushion",
            "COLLATERAL": "4. Collateral — Security Coverage",
            "CONDITIONS": "5. Conditions — External Environment",
        }

        for key in ("CHARACTER", "CAPACITY", "CAPITAL", "COLLATERAL", "CONDITIONS"):
            payload = five_cs_text.get(key)
            if payload is None:
                continue

            # Accept either a plain string or the FiveCsWriter result dict
            text = payload if isinstance(payload, str) else payload.get("text", "")
            if not text:
                continue

            _add_subheading(doc, section_labels.get(key, key))
            _spacer(doc)

            # Split into paragraphs (LLM text may contain blank-line-separated paras)
            for chunk in text.split("\n\n"):
                chunk = chunk.strip()
                if not chunk:
                    continue
                # Strip any markdown heading lines the LLM may have produced
                if chunk.startswith("#"):
                    chunk = chunk.lstrip("#").strip()
                p = doc.add_paragraph(chunk)
                p.paragraph_format.space_after = Pt(6)

            _spacer(doc, 6)

    # ==================================================================
    # Section 7 — Risk Score & SHAP
    # ==================================================================

    def _risk_score_section(
        self,
        doc:            Any,
        scoring_result: dict[str, Any],
    ) -> None:
        _add_navy_heading(doc, "Risk Score & Model Explanation")
        _spacer(doc)

        risk_score   = scoring_result.get("risk_score")
        risk_band    = _safe_str(scoring_result.get("risk_band"))
        default_prob = scoring_result.get("default_probability")
        band_colour  = _BAND_COLOUR.get(risk_band, _BLACK)
        shap         = scoring_result.get("shap_explanations") or {}

        # Score summary 2-col table
        score_rows = [
            ("LightGBM Risk Score",    f"{risk_score:.2f} / 10.00" if risk_score is not None else "N/A"),
            ("Probability of Default", f"{default_prob:.2%}"       if default_prob is not None else "N/A"),
            ("Risk Band",              risk_band),
        ]
        score_tbl = _make_table(doc, len(score_rows), 2, col_widths=[2.8, self._pw - 2.8])
        for i, (lbl, val) in enumerate(score_rows):
            bg = _GRAY if i % 2 == 0 else None
            if bg:
                _set_row_bg(score_tbl.rows[i], bg)
            _cell_text(score_tbl.rows[i].cells[0], lbl, bold=True, font_size=10)
            colour = band_colour if lbl == "Risk Band" else _BLACK
            _cell_text(score_tbl.rows[i].cells[1], val, bold=(lbl == "Risk Band"),
                       colour=colour, font_size=10)

        _spacer(doc, 10)
        _add_subheading(doc, "Top Risk Drivers (SHAP Analysis)")
        _spacer(doc)

        top_risk = shap.get("top_risk_factors") or []
        top_pos  = shap.get("top_positive_factors") or []

        # Risk factors table
        if top_risk:
            rf_tbl = _make_table(doc, len(top_risk) + 1, 3,
                                 col_widths=[2.8, 2.5, 0.97])
            _set_row_bg(rf_tbl.rows[0], _NAVY)
            for j, hdr in enumerate(["Risk Factor", "Description", "SHAP Value"]):
                _cell_text(rf_tbl.rows[0].cells[j], hdr,
                           bold=True, colour=_WHITE, font_size=10)
            for i, factor in enumerate(top_risk[:5]):
                row = rf_tbl.rows[i + 1]
                bg  = _GRAY if i % 2 == 0 else None
                if bg:
                    _set_row_bg(row, bg)
                _cell_text(row.cells[0],
                           factor.get("feature_name", "").replace("_", " ").title(),
                           font_size=10, colour=_RED)
                _cell_text(row.cells[1],
                           _safe_str(factor.get("human_readable_name")),
                           font_size=10)
                sv = factor.get("shap_value")
                sv_str = f"{sv:+.4f}" if sv is not None else "N/A"
                _cell_text(row.cells[2], sv_str, font_size=10,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           colour=_RED if (sv is not None and sv > 0) else _GREEN)

        _spacer(doc, 8)

        if top_pos:
            _add_subheading(doc, "Top Protective Factors (SHAP Analysis)")
            _spacer(doc)
            pp_tbl = _make_table(doc, len(top_pos) + 1, 3,
                                 col_widths=[2.8, 2.5, 0.97])
            _set_row_bg(pp_tbl.rows[0], _NAVY)
            for j, hdr in enumerate(["Protective Factor", "Description", "SHAP Value"]):
                _cell_text(pp_tbl.rows[0].cells[j], hdr,
                           bold=True, colour=_WHITE, font_size=10)
            for i, factor in enumerate(top_pos[:3]):
                row = pp_tbl.rows[i + 1]
                bg  = _GRAY if i % 2 == 0 else None
                if bg:
                    _set_row_bg(row, bg)
                _cell_text(row.cells[0],
                           factor.get("feature_name", "").replace("_", " ").title(),
                           font_size=10, colour=_GREEN)
                _cell_text(row.cells[1],
                           _safe_str(factor.get("human_readable_name")),
                           font_size=10)
                sv = factor.get("shap_value")
                sv_str = f"{sv:+.4f}" if sv is not None else "N/A"
                _cell_text(row.cells[2], sv_str, font_size=10,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           colour=_GREEN if (sv is not None and sv < 0) else _RED)

    # ==================================================================
    # Section 8 — Recommendation
    # ==================================================================

    def _recommendation(
        self,
        doc:            Any,
        company_data:   dict[str, Any],
        scoring_result: dict[str, Any],
    ) -> None:
        _add_navy_heading(doc, "Recommendation")
        _spacer(doc)

        decision    = _safe_str(company_data.get("decision"), "PENDING").upper()
        rec_amount  = _safe_str(company_data.get("recommended_amount"))
        rate        = _safe_str(company_data.get("interest_rate"))
        tenure      = _safe_str(company_data.get("tenure"))
        rationale   = _safe_str(
            company_data.get("decision_rationale"),
            "Subject to satisfactory legal, technical, and financial due-diligence."
        )

        # Coloured recommendation box using a single-cell table
        box_tbl = _make_table(doc, 1, 1, col_widths=[self._pw])
        cell = box_tbl.rows[0].cells[0]

        # Background: green for approve, red for reject, amber for conditional
        box_bg = {
            "APPROVE":     RGBColor(0xE2, 0xEF, 0xDA),   # light green
            "REJECT":      RGBColor(0xFF, 0xE6, 0xE6),   # light red
            "CONDITIONAL": RGBColor(0xFF, 0xF2, 0xCC),   # light amber
        }.get(decision.split(":")[0].strip(), _LIGHT_BLUE)
        _set_cell_bg(cell, box_bg)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Decision line
        dec_run = para.add_run(f"DECISION: {decision}\n")
        dec_run.bold           = True
        dec_run.font.size      = Pt(12)
        dec_colour = _DECISION_COLOUR.get(decision.split(":")[0].strip(), _BLACK)
        dec_run.font.color.rgb = dec_colour

        # Terms
        terms_run = para.add_run(
            f"Recommended Amount: {rec_amount}   |   "
            f"Rate: {rate}   |   Tenure: {tenure}\n\n"
        )
        terms_run.bold      = True
        terms_run.font.size = Pt(10)

        # Rationale (3 sentences)
        rat_run = para.add_run(rationale)
        rat_run.font.size = Pt(10)

        _spacer(doc)

        # Conditions precedent (if any)
        conditions = company_data.get("conditions_precedent") or []
        if conditions:
            _add_subheading(doc, "Conditions Precedent to Disbursement")
            _spacer(doc)
            for cond in conditions:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(cond)).font.size = Pt(10)

    # ==================================================================
    # Section 9 — Early Warning Indicators
    # ==================================================================

    def _early_warning_indicators(
        self,
        doc:          Any,
        company_data: dict[str, Any],
    ) -> None:
        _add_navy_heading(doc, "Early Warning Indicators & Monitoring Triggers")
        _spacer(doc)

        intro = doc.add_paragraph(
            "The following covenant and monitoring triggers have been pre-agreed "
            "with the borrower. Breach of any trigger will initiate a review by "
            "the credit committee within 15 working days."
        )
        intro.paragraph_format.space_after = Pt(8)

        ewi = company_data.get("ewi_triggers") or _DEFAULT_EWI

        for trigger in ewi:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(trigger))
            run.font.size = Pt(10)
            # Highlight any that mention 'CRITICAL' or 'fraud' in red
            if any(w in trigger.lower() for w in ("fraud", "critical", "wilful", "nclt", "ibc", "insolvenc")):
                run.font.color.rgb = _RED

        _spacer(doc, 12)

        # Sign-off footer
        footer_tbl = _make_table(doc, 2, 3,
                                 col_widths=[self._pw / 3] * 3)
        _set_row_bg(footer_tbl.rows[0], _NAVY)
        for j, hdr in enumerate(["Prepared by", "Reviewed by", "Approved by"]):
            _cell_text(footer_tbl.rows[0].cells[j], hdr,
                       bold=True, colour=_WHITE, font_size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for j in range(3):
            cell = footer_tbl.rows[1].cells[j]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run("\n\n\n_________________________\nSignature & Date")
            run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Module-level convenience wrapper
# ---------------------------------------------------------------------------

def generate_cam(
    company_data:    dict[str, Any],
    scoring_result:  dict[str, Any],
    research_report: dict[str, Any],
    five_cs_text:    dict[str, Any],
    output_path:     str | Path = "outputs/credit_appraisal_memo.docx",
) -> Path:
    """One-liner: generate the CAM and return the saved path."""
    return CAMGenerator().generate_cam(
        company_data    = company_data,
        scoring_result  = scoring_result,
        research_report = research_report,
        five_cs_text    = five_cs_text,
        output_path     = output_path,
    )


# ---------------------------------------------------------------------------
# CLI smoke-test
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys as _sys

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(name)s — %(message)s",
    )

    # ── Synthetic demo payload ─────────────────────────────────────────
    company_data = {
        "name":                  "Acme Manufacturing Private Limited",
        "cin":                   "U28999MH2009PTC123456",
        "incorporation_date":    "12 March 2009",
        "registered_office":     "Chakan Industrial Area, Pune, Maharashtra 410501",
        "business_sector":       "Auto-Components Manufacturing",
        "constitution":          "Private Limited Company",
        "business_description":  (
            "Acme Manufacturing Private Limited is a Pune-based manufacturer of "
            "precision auto-components supplying to OEMs in the passenger vehicle "
            "and two-wheeler segments. The company holds three OEM certifications "
            "and operates from its own factory premises spread over 2.5 acres."
        ),
        "directors": [
            {"name": "Rajesh Mehta",   "din": "06781234", "designation": "Managing Director"},
            {"name": "Supriya Mehta",  "din": "06789876", "designation": "Whole-Time Director"},
            {"name": "CA K P Sharma",  "din": "07654321", "designation": "Independent Director"},
        ],
        "loan_amount_requested": "₹50 Crore",
        "recommended_amount":    "₹45 Crore",
        "interest_rate":         "12.50% p.a. (floating, quarterly reset)",
        "tenure":                "60 months (including 6-month moratorium)",
        "decision":              "CONDITIONAL: Subject to charge registration and auditor NOC",
        "decision_rationale":    (
            "The borrower demonstrates adequate debt servicing capacity with a DSCR of 1.42 "
            "and consistent revenue growth of 12.4% CAGR over three years. "
            "The collateral coverage of 1.7x provides reasonable security buffer, though "
            "subordination of the existing SBI charge must be completed prior to disbursement. "
            "The credit committee recommends conditional approval pending fulfilment of "
            "pre-disbursement conditions set out herein."
        ),
        "conditions_precedent": [
            "Registered mortgage of industrial property in favour of the bank",
            "Subordination agreement / NOC from State Bank of India for existing charge",
            "Submission of latest three months bank statements from all lenders",
            "Auditor certificate confirming no qualifications in the current year",
            "Personal guarantee documentation from Rajesh Mehta executed and registered",
        ],
        "financials_3yr": [
            {"year": "FY 2023-24", "revenue": 118.4, "ebitda": 16.9, "pat": 6.1,
             "ebitda_margin_pct": 14.3, "pat_margin_pct": 5.2,
             "de_ratio": 1.51, "current_ratio": 1.05, "dscr": 1.19,
             "revenue_growth_pct": 8.2},
            {"year": "FY 2024-25", "revenue": 131.6, "ebitda": 19.1, "pat": 7.4,
             "ebitda_margin_pct": 14.5, "pat_margin_pct": 5.6,
             "de_ratio": 1.44, "current_ratio": 1.12, "dscr": 1.31,
             "revenue_growth_pct": 11.1},
            {"year": "FY 2025-26", "revenue": 148.0, "ebitda": 22.1, "pat": 8.6,
             "ebitda_margin_pct": 14.9, "pat_margin_pct": 5.8,
             "de_ratio": 1.38, "current_ratio": 1.18, "dscr": 1.42,
             "revenue_growth_pct": 12.4},
        ],
        "gst_findings": {
            "health_score":         7.8,
            "grade":               "B+",
            "itc_gap_pct":          8.4,
            "turnover_consistency": "CLEAN",
            "filing_regularity":    "REGULAR",
            "fictitious_vendors":   0,
        },
        "bank_findings": {
            "avg_monthly_balance": "₹1.8 Cr",
            "debit_credit_ratio":   0.92,
            "bounce_count":         1,
            "upi_concentration":    18.5,
            "cash_stress_flag":    "LOW",
            "revenue_inflation_flag": "CLEAR",
        },
        "ews_flags": {
            "gst_itc_fraud_risk":     "LOW",
            "circular_trading_risk":  "CLEAR",
            "revenue_inflation_risk": "CLEAR",
            "cash_stress_risk":       "LOW",
            "documentation_risk":     "CLEAR",
            "auditor_concern_risk":   "CLEAR",
            "director_risk":          "CLEAR",
            "compliance_risk":        "CLEAR",
        },
    }

    scoring_result = {
        "default_probability": 0.18,
        "risk_score":           8.2,
        "risk_band":           "PRIME",
        "raw_lgbm_proba":       0.18,
        "shap_explanations": {
            "method": "shap_tree_explainer",
            "top_risk_factors": [
                {"feature_name": "debt_to_equity",       "human_readable_name": "Debt-to-Equity Ratio",        "shap_value": +0.142, "direction": "INCREASES_DEFAULT_RISK"},
                {"feature_name": "bounce_count",         "human_readable_name": "Cheque / ECS Bounce Count",   "shap_value": +0.031, "direction": "INCREASES_DEFAULT_RISK"},
                {"feature_name": "itc_gap_pct",          "human_readable_name": "ITC Gap vs GSTR-2A (%)",      "shap_value": +0.018, "direction": "INCREASES_DEFAULT_RISK"},
            ],
            "top_positive_factors": [
                {"feature_name": "current_ratio",        "human_readable_name": "Current Ratio",               "shap_value": -0.210, "direction": "DECREASES_DEFAULT_RISK"},
                {"feature_name": "dscr",                 "human_readable_name": "Debt Service Coverage Ratio", "shap_value": -0.189, "direction": "DECREASES_DEFAULT_RISK"},
                {"feature_name": "news_risk_score",      "human_readable_name": "News Risk Score (0–10)",       "shap_value": -0.054, "direction": "DECREASES_DEFAULT_RISK"},
            ],
        },
    }

    research_report = {
        "news_summary":                  "No adverse media coverage detected in the last 24 months.",
        "promoter_risk_flag":            "LOW: No RBI wilful-defaulter matches found.",
        "regulatory_compliance_summary": "MCA filings current; no SEBI or RBI notices outstanding.",
        "key_red_flags":                 ["Single existing charge held by SBI (₹22 Cr unsatisfied)"],
        "positive_signals":              ["Consistent revenue growth", "Clean RBI defaulter record"],
        "overall_external_risk_score":    2.5,
        "recommended_action":            "PROCEED: External intelligence broadly clean.",
    }

    five_cs_text = {
        "CHARACTER":  (
            "The promoters of Acme Manufacturing Private Limited, Rajesh Mehta and Supriya Mehta, "
            "demonstrate a satisfactory track record with fifteen years of industry experience. "
            "No wilful default flag is recorded against the company or its directors. "
            "The promoters maintain good standing with banking institutions and exhibit sound governance "
            "through structured succession planning and second-generation management involvement."
        ),
        "CAPACITY":   (
            "The borrower demonstrates adequate debt servicing capacity supported by positive "
            "operating cash generation of ₹14 Crore and a DSCR of 1.42. "
            "The current ratio of 1.18 indicates adequate liquidity. "
            "The interest coverage ratio of 3.1 comfortably covers interest obligations and "
            "supports the proposed credit facility repayment schedule."
        ),
        "CAPITAL":    (
            "The net worth stands at ₹42 Crore with tangible net worth of ₹39 Crore, "
            "representing a moderate equity cushion. The debt-to-equity ratio of 1.38 "
            "is within acceptable parameters for the sector. Promoter shareholding of 74% "
            "demonstrates strong alignment of interests with lenders."
        ),
        "COLLATERAL": (
            "Primary security comprises industrial land and factory building in Chakan Industrial Area, "
            "Pune, valued at ₹85 Crore as of January 2026. The security coverage ratio of 1.7 "
            "provides adequate buffer. One existing charge of ₹22 Crore (SBI) requires "
            "subordination prior to disbursement. Personal guarantee of Rajesh Mehta (NW ₹28 Cr) "
            "is offered as secondary security."
        ),
        "CONDITIONS": (
            "India's auto-components sector is growing at 8.5% against GDP expansion of 6.9%. "
            "The stable interest rate environment (repo 6.5%) is supportive. Key risks include "
            "commodity input-cost inflation impacting margins and medium-term EV transition risk. "
            "PLI scheme benefits provide structural regulatory support."
        ),
    }

    out_path = _PROJECT_ROOT / "outputs" / "ACME_CAM_DEMO.docx"
    gen  = CAMGenerator()
    saved = gen.generate_cam(
        company_data    = company_data,
        scoring_result  = scoring_result,
        research_report = research_report,
        five_cs_text    = five_cs_text,
        output_path     = out_path,
    )
    print(f"\n✓ CAM saved → {saved}")
    print(f"  File size: {saved.stat().st_size:,} bytes\n")
